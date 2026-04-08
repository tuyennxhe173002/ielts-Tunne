from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
DOCS = [
    ROOT / "docs" / "system-architecture.md",
    ROOT / "docs" / "database-schema.md",
    ROOT / "docs" / "permission-model.md",
]
OUTPUT = ROOT / "docs" / "lms-architecture-package.docx"


def main() -> None:
    doc = Document()
    doc.add_heading("LMS Architecture Package", 0)

    for file_path in DOCS:
        if file_path != DOCS[0]:
            doc.add_page_break()
        render_markdown(doc, file_path.read_text(encoding="utf-8").splitlines())

    doc.save(str(OUTPUT))


def render_markdown(doc: Document, lines: list[str]) -> None:
    for line in lines:
        text = line.rstrip()
        if not text:
            doc.add_paragraph("")
            continue
        if text.startswith("### "):
            doc.add_heading(text[4:], level=3)
            continue
        if text.startswith("## "):
            doc.add_heading(text[3:], level=2)
            continue
        if text.startswith("# "):
            doc.add_heading(text[2:], level=1)
            continue
        if text.startswith("- "):
            doc.add_paragraph(text[2:], style="List Bullet")
            continue
        if text.startswith("| ") or text.startswith("```"):
            doc.add_paragraph(text)
            continue
        numbered = parse_numbered(text)
        if numbered:
            doc.add_paragraph(numbered, style="List Number")
            continue
        doc.add_paragraph(text)


def parse_numbered(text: str) -> str | None:
    if ". " not in text:
        return None
    prefix, content = text.split(". ", 1)
    return content if prefix.isdigit() else None


if __name__ == "__main__":
    main()
